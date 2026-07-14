#!/usr/bin/env python3
"""
Professor System Prompt Builder.
Generates a structured system prompt for LLM-based persona simulation from paper evidence.
"""

import argparse
import json
import os
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

import requests
from docx import Document
from docx.oxml.ns import qn


DEFAULT_BASE_URL = "http://58.210.177.113:8888/v1"
DEFAULT_MODEL = "mimo-v2-flash"
REQUEST_TIMEOUT = 300
DEFAULT_MAX_TOKENS = 2200
STYLE_DIMENSIONS = [
    ("Problem-driven", "Method-driven"),
    ("Theory-oriented", "Application-oriented"),
    ("System-integration", "Modular-breakthrough"),
    ("Conservative-iteration", "Radical-exploration"),
    ("Empirical-induction", "Mechanistic-explanation"),
    ("Engineering-evidence", "Formal-modeling"),
]


@dataclass
class PaperDigest:
    title: str
    abstract: str
    keywords: str
    conclusion: str
    representativeness_score: float = 0.0
    bucket: str = "unknown"
    matched_topic: str = ""


def normalize_text(text: str) -> str:
    return unicodedata.normalize("NFKC", text or "").strip()


def cut_text(text: str, max_len: int) -> str:
    text = normalize_text(text)
    if len(text) <= max_len:
        return text
    return text[:max_len].rstrip() + "..."


def sanitize_filename_component(name: str) -> str:
    name = normalize_text(name)
    name = re.sub(r"[<>:\"/\\\\|?*]", "_", name)
    name = re.sub(r"\s+", " ", name).strip().strip(".")
    return name or "Professor"


def apply_document_fonts(doc: Document) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def parse_author_name(paragraphs: List[str], fallback_name: str) -> str:
    for line in paragraphs[:8]:
        t = normalize_text(line)
        if not t:
            continue
        if "-" in t:
            maybe = t.split("-")[-1].strip()
            if maybe and len(maybe) <= 80:
                return maybe
        if "—" in t:
            maybe = t.split("—")[-1].strip()
            if maybe and len(maybe) <= 80:
                return maybe
    return fallback_name


def parse_papers_from_docx(docx_path: Path) -> tuple[str, List[PaperDigest]]:
    doc = Document(str(docx_path))
    paragraphs = [normalize_text(p.text) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]

    inferred_name = docx_path.stem
    inferred_name = re.sub(r"(report|digest|literature.review)$", "", inferred_name, flags=re.IGNORECASE).strip()
    author_name = parse_author_name(paragraphs, inferred_name or "Unknown Professor")

    papers: List[PaperDigest] = []
    current_title: Optional[str] = None
    abstract = ""
    keywords = ""
    conclusion = ""
    mode = None

    def flush_current():
        nonlocal current_title, abstract, keywords, conclusion, mode
        if current_title:
            papers.append(
                PaperDigest(
                    title=cut_text(current_title, 300),
                    abstract=cut_text(abstract, 1800),
                    keywords=cut_text(keywords, 500),
                    conclusion=cut_text(conclusion, 1800),
                )
            )
        current_title = None
        abstract = ""
        keywords = ""
        conclusion = ""
        mode = None

    heading_pattern = re.compile(r"^\d+\s*[\.、]\s*(.+)$")

    for line in paragraphs:
        m = heading_pattern.match(line)
        if m:
            flush_current()
            current_title = m.group(1).strip()
            continue

        if line.startswith("Abstract") or line.startswith("摘要"):
            mode = "abstract"
            continue
        if line.startswith("Keywords") or line.startswith("关键词") or line.startswith("关键字"):
            mode = "keywords"
            continue
        if line.startswith("Conclusion") or line.startswith("结论"):
            mode = "conclusion"
            continue

        if current_title is None:
            continue

        if mode == "abstract":
            abstract = f"{abstract}\n{line}".strip()
        elif mode == "keywords":
            keywords = f"{keywords}; {line}".strip("; ")
        elif mode == "conclusion":
            conclusion = f"{conclusion}\n{line}".strip()

    flush_current()
    return author_name, papers


def parse_papers_from_json(json_path: Path) -> tuple[str, List[PaperDigest]]:
    with open(json_path, "r", encoding="utf-8") as f:
        obj = json.load(f)

    author_name = normalize_text(obj.get("author_name") or json_path.stem or "Unknown Professor")
    papers: List[PaperDigest] = []
    for item in obj.get("papers") or []:
        papers.append(
            PaperDigest(
                title=cut_text(str(item.get("title") or ""), 300),
                abstract=cut_text(str(item.get("abstract") or ""), 1800),
                keywords=cut_text("; ".join(item.get("keywords") or []) if isinstance(item.get("keywords"), list) else str(item.get("keywords") or ""), 500),
                conclusion=cut_text(str(item.get("conclusion") or ""), 1800),
                representativeness_score=float(item.get("representativeness_score") or 0.0),
                bucket=str(item.get("bucket") or "unknown"),
                matched_topic=str(item.get("matched_topic") or ""),
            )
        )

    papers.sort(key=lambda x: x.representativeness_score, reverse=True)
    return author_name, papers


def summarize_mainline_graph(mainline_graph_path: Optional[Path]) -> str:
    if not mainline_graph_path:
        return "No mainline graph provided."
    if not mainline_graph_path.exists():
        return "Mainline graph file not found."
    try:
        with open(mainline_graph_path, "r", encoding="utf-8") as f:
            obj = json.load(f)
    except Exception:
        return "Failed to read mainline graph file."

    nodes = obj.get("nodes") or []
    timeline = obj.get("timeline") or []
    if not nodes:
        return "Mainline graph has no valid topic nodes."

    nodes_sorted = sorted(nodes, key=lambda x: float(x.get("weight") or 0.0), reverse=True)
    top_nodes = nodes_sorted[:6]
    node_lines = []
    for n in top_nodes:
        node_lines.append(
            f"- Topic: {n.get('topic')}, Weight: {float(n.get('weight') or 0.0):.3f}, "
            f"Period: {n.get('first_year') or 'Unknown'}-{n.get('last_year') or 'Unknown'}, "
            f"Evidence: {int(n.get('evidence_count') or 0)}"
        )

    timeline_lines = []
    for t in timeline[:8]:
        topics = ", ".join(t.get("top_topics") or [])
        timeline_lines.append(f"- {t.get('year')}: {topics}")

    return "\n".join(
        [
            "Mainline topic nodes:",
            *node_lines,
            "Timeline evolution:",
            *(timeline_lines or ["- None"]),
        ]
    )


def summarize_external_sources(sources_path: Optional[Path], max_items: int = 24) -> str:
    if not sources_path:
        return "No external sources (interviews/talks/homepages) provided."
    if not sources_path.exists():
        return "External sources file not found."
    try:
        raw_lines = sources_path.read_text(encoding="utf-8", errors="ignore").splitlines()
    except Exception:
        return "Failed to read external sources file."

    items: List[str] = []
    seen = set()
    for line in raw_lines:
        t = normalize_text(line)
        if (not t) or t.startswith("#"):
            continue
        if t in seen:
            continue
        seen.add(t)
        items.append(t)
        if len(items) >= max(1, int(max_items)):
            break

    if not items:
        return "External sources file is empty or has no valid content."
    return "\n".join([f"- {cut_text(x, 240)}" for x in items])


def evaluate_evidence_sufficiency(papers: List[PaperDigest]) -> tuple[bool, List[str]]:
    warnings: List[str] = []
    core_count = len([p for p in papers if p.bucket == "core"])
    if core_count < 3:
        warnings.append(f"Only {core_count} core papers — profile confidence is low.")
    if len(papers) < 6:
        warnings.append(f"Only {len(papers)} total papers — evidence coverage may be insufficient.")
    return len(warnings) == 0, warnings


def build_user_prompt(
    author_name: str,
    papers: List[PaperDigest],
    mainline_summary: str = "",
    external_sources_summary: str = "",
    core_limit: int = 12,
    var_limit: int = 8,
    fallback_core_limit: int = 8,
    evidence_warnings: Optional[List[str]] = None,
) -> str:
    core_papers = [p for p in papers if p.bucket == "core"][: max(1, int(core_limit))]
    peripheral_papers = [p for p in papers if p.bucket in {"peripheral", "student-led"}][: max(0, int(var_limit))]
    if not core_papers:
        core_papers = papers[: max(1, int(fallback_core_limit))]

    core_lines = []
    for i, p in enumerate(core_papers, 1):
        core_lines.append(
            "\n".join(
                [
                    f"[Core-{i}] Title: {p.title}",
                    f"[Core-{i}] Topic Match: {p.matched_topic or 'Unlabeled'}",
                    f"[Core-{i}] Representativeness: {p.representativeness_score:.3f}",
                    f"[Core-{i}] Abstract: {p.abstract or 'Not available'}",
                    f"[Core-{i}] Keywords: {p.keywords or 'Not available'}",
                    f"[Core-{i}] Conclusion: {p.conclusion or 'Not available'}",
                ]
            )
        )

    peripheral_lines = []
    for i, p in enumerate(peripheral_papers, 1):
        peripheral_lines.append(
            "\n".join(
                [
                    f"[Var-{i}] Title: {p.title}",
                    f"[Var-{i}] Bucket: {p.bucket}",
                    f"[Var-{i}] Topic Match: {p.matched_topic or 'Unlabeled'}",
                    f"[Var-{i}] Abstract: {p.abstract or 'Not available'}",
                ]
            )
        )

    dimensions_text = "\n".join([f"- {left} <-> {right}" for left, right in STYLE_DIMENSIONS])
    core_evidence = "\n\n".join(core_lines) if core_lines else "None"
    var_evidence = "\n\n".join(peripheral_lines) if peripheral_lines else "None"
    mainline_summary = mainline_summary or "No mainline graph provided."
    external_sources_summary = external_sources_summary or "No external sources provided."
    evidence_warning_text = "\n".join([f"- {w}" for w in (evidence_warnings or [])]) or "- None"

    return f"""Based on the given evidence, generate a System Prompt for professor {author_name} that can be directly used in an AI Agent.

Hard requirements:
1) Use second person (you/your).
2) Output must be the system prompt body only — no preamble, postscript, or explanation.
3) Sections must be clear and must include these top-level headings:
- Role Definition
- Primary Research Areas and Representative Work
- Academic Style and Research Habits
- Research Philosophy and Decision Principles
- Communication Style and Mentoring Preferences
- Behavioral Boundaries and Taboos
4) In "Primary Research Areas and Representative Work", first present "Research Mainline and Timeline Evolution", then representative work.
5) "Academic Style and Research Habits" must be layered:
Layer 1: Academic style labels (comparable dimensions)
Layer 2: Personalized academic style description
6) For Layer 1, label each dimension below with evidence-based justification for which side you lean toward:
{dimensions_text}
7) Style characterization should prioritize Core evidence; Var evidence is only for "group diversity notes" — avoid mistaking student styles for the advisor's stable style.
8) Each section must contain actionable rules, not vague value judgments.
9) Do not fabricate directions not present in the evidence; reasonable inference is allowed but must be consistent with evidence.
10) Output in English.
11) If evidence is insufficient, explicitly mark "Low-confidence hypothesis" — do not present guesses as established facts.
12) In "Communication Style and Mentoring Preferences" and "Behavioral Boundaries and Taboos", prioritize external sources (interviews/talks/homepages); papers serve only as supporting evidence.
13) When external sources are insufficient, explicitly mark low confidence — do not equate technical writing style with mentoring/communication style.

Evidence sufficiency warnings:
{evidence_warning_text}

Research mainline graph summary:
{mainline_summary}

External source evidence (interviews/talks/homepages/news — prioritize for communication/mentoring/boundaries):
{external_sources_summary}

Core evidence (for stable style and mainline):
{core_evidence}

Var evidence (for group diversity notes only):
{var_evidence}
"""


def call_llm(
    base_url: str,
    api_key: str,
    model: str,
    user_prompt: str,
    max_tokens: int,
    thinking: bool,
    proxies: Optional[dict] = None,
) -> tuple[str, str]:
    endpoint = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a senior academic profiling assistant, skilled at distilling actionable, distinguishable professor persona system prompts from paper evidence.",
            },
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.2,
        "max_tokens": int(max_tokens),
        "thinking": bool(thinking),
    }

    session = requests.Session()
    session.trust_env = False
    try:
        resp = session.post(
            endpoint,
            headers=headers,
            json=payload,
            timeout=REQUEST_TIMEOUT,
            proxies=proxies,
        )
    except requests.exceptions.ReadTimeout:
        if proxies:
            resp = session.post(
                endpoint,
                headers=headers,
                json=payload,
                timeout=REQUEST_TIMEOUT,
                proxies={},
            )
        else:
            raise
    resp.raise_for_status()
    data = resp.json()
    choice0 = (data.get("choices") or [{}])[0]
    message = choice0.get("message") or {}
    content = message.get("content")
    reasoning_content = message.get("reasoning_content")

    finish_reason = str(choice0.get("finish_reason") or "").strip().lower()

    if isinstance(content, str) and content.strip():
        return normalize_text(content), finish_reason
    if isinstance(reasoning_content, str) and reasoning_content.strip():
        return normalize_text(reasoning_content), finish_reason

    # Some providers return content as a list of chunks.
    if isinstance(content, list):
        chunks = []
        for item in content:
            if isinstance(item, dict):
                txt = item.get("text") or item.get("content")
                if isinstance(txt, str) and txt.strip():
                    chunks.append(txt.strip())
            elif isinstance(item, str) and item.strip():
                chunks.append(item.strip())
        if chunks:
            return normalize_text("\n".join(chunks)), finish_reason

    raise RuntimeError("No usable text found in model output (message.content/reasoning_content both empty).")


def has_required_sections(text: str) -> bool:
    required = [
        "Role Definition",
        "Primary Research Areas and Representative Work",
        "Academic Style and Research Habits",
        "Layer 1: Academic Style Labels",
        "Layer 2: Personalized Academic Style Description",
        "Research Philosophy and Decision Principles",
        "Communication Style and Mentoring Preferences",
        "Behavioral Boundaries and Taboos",
    ]
    return all(k in text for k in required)


def looks_truncated(text: str, finish_reason: str) -> bool:
    if finish_reason == "length":
        return True
    text = (text or "").strip()
    if not text:
        return True
    if text[-1] in {",", ";", ":", "-", "(", "["}:
        return True
    return False


def save_system_prompt_docx(author_name: str, system_prompt: str, output_path: Path) -> None:
    doc = Document()
    apply_document_fonts(doc)
    doc.add_heading(f"{author_name}'s System Prompt", level=1)
    for line in system_prompt.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate professor system prompt from paper extraction results.")
    parser.add_argument("--input", default="", help="Paper extraction Word file path (with title/abstract/keywords/conclusion)")
    parser.add_argument("--input-json", default="", help="Optional: paper scoring JSON (paper_digest_scored.json)")
    parser.add_argument("--mainline-graph", default="", help="Optional: mainline graph JSON (mainline_graph.json)")
    parser.add_argument(
        "--sources-file",
        default="",
        help="Optional: sources text file (e.g. sources.txt, with homepage/interview/talk links or summaries)",
    )
    parser.add_argument("--author", default="", help="Optional: override auto-detected author name")
    parser.add_argument("--output", default="", help="Optional: output filename, default \"XXX's System Prompt v2.docx\"")
    parser.add_argument("--model", default=DEFAULT_MODEL, help=f"LLM model name, default {DEFAULT_MODEL}")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL, help=f"LLM API base URL, default {DEFAULT_BASE_URL}")
    parser.add_argument("--api-key", default="", help="API Key, falls back to OPENAI_API_KEY env var")
    parser.add_argument(
        "--max-tokens",
        type=int,
        default=DEFAULT_MAX_TOKENS,
        help=f"Max generation tokens, default {DEFAULT_MAX_TOKENS}",
    )
    parser.add_argument(
        "--thinking",
        default="false",
        help="Enable thinking mode (true/false), default false",
    )
    parser.add_argument("--core-limit", type=int, default=12, help="Max core evidence papers, default 12")
    parser.add_argument("--var-limit", type=int, default=8, help="Max var evidence papers, default 8")
    parser.add_argument(
        "--fallback-core-limit",
        type=int,
        default=8,
        help="When no core papers, take top N by score, default 8",
    )
    parser.add_argument(
        "--strict-evidence",
        default="false",
        help="Fail when evidence is insufficient (true/false), default false",
    )
    parser.add_argument("--http-proxy", default="", help="HTTP proxy, e.g. http://127.0.0.1:7890")
    parser.add_argument("--https-proxy", default="", help="HTTPS proxy, e.g. http://127.0.0.1:7890")
    args = parser.parse_args()

    input_path = None
    if args.input_json.strip():
        input_path = Path(args.input_json).expanduser().resolve()
        if not input_path.exists():
            raise FileNotFoundError(f"Input JSON not found: {input_path}")
        author_name, papers = parse_papers_from_json(input_path)
    elif args.input.strip():
        input_path = Path(args.input).expanduser().resolve()
        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        author_name, papers = parse_papers_from_docx(input_path)
    else:
        raise RuntimeError("Please provide --input or --input-json.")

    if args.author.strip():
        author_name = args.author.strip()

    if not papers:
        raise RuntimeError("No paper entries parsed from input. Check file format.")

    api_key = (args.api_key or os.getenv("OPENAI_API_KEY") or "").strip()
    if not api_key:
        raise RuntimeError("Missing API Key. Provide via --api-key or OPENAI_API_KEY environment variable.")

    proxies = {}
    if args.http_proxy.strip():
        proxies["http"] = args.http_proxy.strip()
    if args.https_proxy.strip():
        proxies["https"] = args.https_proxy.strip()
    proxies = proxies or None

    mainline_path = Path(args.mainline_graph).expanduser().resolve() if args.mainline_graph.strip() else None
    mainline_summary = summarize_mainline_graph(mainline_path)
    sources_path = Path(args.sources_file).expanduser().resolve() if args.sources_file.strip() else None
    external_sources_summary = summarize_external_sources(sources_path)
    sufficient, evidence_warnings = evaluate_evidence_sufficiency(papers)
    strict_evidence = str(args.strict_evidence).strip().lower() in {"1", "true", "yes", "y", "on"}
    if strict_evidence and (not sufficient):
        raise RuntimeError("Insufficient evidence, aborted by --strict-evidence. Details: " + " | ".join(evidence_warnings))
    user_prompt = build_user_prompt(
        author_name,
        papers,
        mainline_summary=mainline_summary,
        external_sources_summary=external_sources_summary,
        core_limit=max(1, int(args.core_limit)),
        var_limit=max(0, int(args.var_limit)),
        fallback_core_limit=max(1, int(args.fallback_core_limit)),
        evidence_warnings=evidence_warnings,
    )
    thinking = str(args.thinking).strip().lower() in {"1", "true", "yes", "y", "on"}
    system_prompt = ""
    finish_reason = ""
    token_budget = max(600, int(args.max_tokens))
    for _ in range(3):
        system_prompt, finish_reason = call_llm(
            args.base_url,
            api_key,
            args.model,
            user_prompt,
            token_budget,
            thinking,
            proxies=proxies,
        )
        if (not looks_truncated(system_prompt, finish_reason)) and has_required_sections(system_prompt):
            break
        token_budget = min(int(token_budget * 1.8), 6000)

    if args.output.strip():
        output_path = Path(args.output).expanduser()
        if not output_path.is_absolute():
            output_path = input_path.parent / output_path
    else:
        safe_name = sanitize_filename_component(author_name)
        output_path = input_path.parent / f"{safe_name}'s System Prompt v2.docx"

    save_system_prompt_docx(author_name, system_prompt, output_path)
    print(f"Done: {output_path}")


# ============ Callable interface for pipeline ============

def run_single(scholar_name: str, output_dir: str, digest_file: str = None,
               mainline_file: str = None, sources_file: str = None,
               language: str = "en", api_key: str = "",
               base_url: str = DEFAULT_BASE_URL, model: str = DEFAULT_MODEL) -> str:
    """Generate system prompt for a single scholar.

    Args:
        scholar_name: Full name of the scholar.
        output_dir: Output directory path.
        digest_file: Path to scored paper JSON.
        mainline_file: Path to mainline graph JSON.
        sources_file: Path to sources.txt.
        language: Output language ("en" or "zh").
        api_key: LLM API key.
        base_url: LLM API base URL.
        model: LLM model name.

    Returns:
        Path to the generated system prompt docx.
    """
    import csv as csv_mod

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    # Auto-discover paths
    if not digest_file:
        digest_file = str(out / "paper_digest_scored.json")
    if not mainline_file:
        mainline_file = str(out / "mainline_graph.json")
    if not sources_file:
        safe = sanitize_filename_component(scholar_name)
        sources_file = str(out / f"{safe}_sources.txt")

    digest_path = Path(digest_file)
    if not digest_path.exists():
        raise FileNotFoundError(f"Digest file not found: {digest_path}")

    author_name, papers = parse_papers_from_json(digest_path)
    if not papers:
        raise RuntimeError("No papers parsed from digest file.")

    if not api_key:
        api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("Missing API key for prompt generation.")

    mainline_path = Path(mainline_file) if mainline_file and Path(mainline_file).exists() else None
    mainline_summary = summarize_mainline_graph(mainline_path)

    sources_path = Path(sources_file) if sources_file and Path(sources_file).exists() else None
    external_sources_summary = summarize_external_sources(sources_path)

    sufficient, evidence_warnings = evaluate_evidence_sufficiency(papers)

    user_prompt = build_user_prompt(
        author_name,
        papers,
        mainline_summary=mainline_summary,
        external_sources_summary=external_sources_summary,
        evidence_warnings=evidence_warnings,
    )

    system_prompt = ""
    finish_reason = ""
    token_budget = 2200
    for _ in range(3):
        system_prompt, finish_reason = call_llm(
            base_url, api_key, model, user_prompt, token_budget, False, proxies=None,
        )
        if (not looks_truncated(system_prompt, finish_reason)) and has_required_sections(system_prompt):
            break
        token_budget = min(int(token_budget * 1.8), 6000)

    output_path = out / f"{sanitize_filename_component(author_name)}'s System Prompt v2.docx"
    save_system_prompt_docx(author_name, system_prompt, output_path)

    # Also save as markdown
    md_path = out / "system_prompt.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(system_prompt)

    print(f"[professor_system_prompt_builder] Done: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    main()
