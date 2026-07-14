#!/usr/bin/env python3
"""
Scholar paper digest extraction module.
Extracts, scores, and ranks papers from DBLP for a given scholar.
"""

import argparse
import json
import math
import os
import re
import time
import unicodedata
import xml.etree.ElementTree as ET
from collections import Counter
from dataclasses import asdict, dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from scholarly import ProxyGenerator, scholarly

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

# Import shared utilities
from utils import (
    DEFAULT_HEADERS,
    DEFAULT_TIMEOUT,
    canonical_title_key,
    classify_url,
    configure_proxies,
    extract_arxiv_id,
    extract_dblp_pid,
    extract_scholar_user_id,
    get_author_position,
    get_timeout_for_url,
    is_top_n_author,
    normalize_doi,
    normalize_text,
    request_get_with_retry,
    safe_name,
    setup_logging,
    split_authors,
    xml_safe_text,
)

logger = setup_logging(__name__)


# ============ Data classes ============

@dataclass
class PaperResult:
    """Paper result data class."""
    title: str
    year: Optional[int]
    authors: str
    author_position: int
    is_first_author: bool
    is_top_three_author: bool
    is_corresponding_author: bool
    abstract: str
    keywords: List[str]
    conclusion: str
    source_url: Optional[str]
    cited_by_count: int
    authorship_weight: float = 0.0
    mainline_relevance: float = 0.0
    topic_persistence: float = 0.0
    citation_impact: float = 0.0
    external_evidence_bonus: float = 0.0
    self_mention_bonus: float = 0.0
    representativeness_score: float = 0.0
    bucket: str = ""
    matched_topic: str = ""


# ============ Constants ============

OPENALEX_BASE = "https://api.openalex.org/works"
S2_BASE = "https://api.semanticscholar.org/graph/v1/paper"
ARXIV_API = "https://export.arxiv.org/api/query"
MAX_PAPERS = 50
REQUEST_PROXIES = None


# ============ HTTP requests ============

def request_get(url: str, **kwargs) -> Optional[requests.Response]:
    """Send GET request with retry and longer timeout for slow sites."""
    if "timeout" not in kwargs:
        kwargs["timeout"] = get_timeout_for_url(url)
    try:
        return request_get_with_retry(url, max_retries=3, backoff=2.0, **kwargs)
    except (KeyboardInterrupt, SystemExit):
        raise
    except BaseException as exc:
        logger.warning(f"request_get failed for {url}: {type(exc).__name__}: {exc}")
        return None


def fetch_html(url: str) -> Optional[str]:
    """Fetch HTML content from URL."""
    try:
        resp = request_get(
            url,
            proxies=REQUEST_PROXIES,
            headers=DEFAULT_HEADERS,
            allow_redirects=True,
        )
        resp.raise_for_status()
        ctype = (resp.headers.get("content-type") or "").lower()
        if "html" in ctype or "<html" in resp.text[:2000].lower():
            return resp.text
    except Exception:
        return None
    return None


# ============ OpenAlex API ============

def inverted_index_to_text(inv_idx: dict) -> str:
    """Convert OpenAlex inverted index to text."""
    if not inv_idx:
        return ""
    pos_word = {}
    for word, positions in inv_idx.items():
        for p in positions:
            pos_word[p] = word
    words = [pos_word[p] for p in sorted(pos_word.keys())]
    return " ".join(words)


def openalex_abstract_from_doi(doi: str) -> Tuple[Optional[str], List[str]]:
    """Fetch abstract and keywords from OpenAlex API via DOI."""
    doi_norm = normalize_doi(doi)
    if not doi_norm:
        return None, []
    try:
        resp = request_get(
            f"{OPENALEX_BASE}/https://doi.org/{doi_norm}",
            params={"select": "abstract_inverted_index,keywords"},
            proxies=REQUEST_PROXIES,
        )
        resp.raise_for_status()
        obj = resp.json()
        if isinstance(obj, dict) and "results" in obj:
            obj = obj["results"][0] if obj["results"] else {}
        abstract = inverted_index_to_text(obj.get("abstract_inverted_index", {})).strip() or None
        keywords = [k["display_name"] for k in obj.get("keywords", []) if k.get("display_name")]
        return abstract, keywords
    except Exception:
        return None, []


def search_openalex_work(title: str, year: Optional[int]) -> Optional[dict]:
    """Search for paper in OpenAlex."""
    params = {
        "search": title,
        "per-page": 5,
        "select": "id,display_name,publication_year,cited_by_count,abstract_inverted_index,keywords,authorships,open_access,primary_location,best_oa_location,doi,ids",
    }
    try:
        resp = request_get(OPENALEX_BASE, params=params, proxies=REQUEST_PROXIES)
        resp.raise_for_status()
        results = resp.json().get("results", [])
    except Exception:
        return None

    if not results:
        return None

    n_title = normalize_text(title)
    best = None
    best_score = -1
    for w in results:
        w_title = normalize_text(w.get("display_name", ""))
        score = 0
        if w_title == n_title:
            score += 100
        if n_title and (n_title in w_title or w_title in n_title):
            score += 50
        pub_year = w.get("publication_year")
        if year and pub_year and abs(pub_year - year) <= 1:
            score += 20
        if score > best_score:
            best_score = score
            best = w

    return best if best_score >= 30 else None


def is_corresponding_author(author_name: str, work: Optional[dict]) -> bool:
    """Check if author is the corresponding author."""
    if not work:
        return False
    target = normalize_text(author_name)
    for a in work.get("authorships", []):
        disp = normalize_text(a.get("author", {}).get("display_name", ""))
        if target in disp or disp in target:
            if a.get("is_corresponding") is True:
                return True
    return False


# ============ HTML metadata extraction ============

def _tryStructuredMeta(html: str) -> Tuple[Optional[str], List[str]]:
    """Try to extract structured metadata from HTML."""
    soup = BeautifulSoup(html, "html.parser")
    abstract, keywords = None, []

    # DC.meta + citation_abstract
    for name in (
        "citation_abstract", "description", "dc.description", "dc.abstract",
        "og:description", "twitter:description",
    ):
        tag = soup.find("meta", attrs={"name": name}) or soup.find("meta", attrs={"property": name})
        if tag and tag.get("content"):
            abstract = tag["content"].strip()
            break

    # keywords
    for name in ("citation_keywords", "keywords", "dc.keywords", "subject"):
        tag = soup.find("meta", attrs={"name": name}) or soup.find("meta", attrs={"property": name})
        if tag and tag.get("content"):
            raw = tag["content"].strip()
            keywords = [k.strip() for k in re.split(r"[,;]", raw) if k.strip()]
            break

    # JSON-LD
    for script in soup.find_all("script", type="application/ld+json"):
        try:
            obj = json.loads(script.string or "")
            for item in (obj if isinstance(obj, list) else [obj]):
                if item.get("@type") in ("ScholarlyArticle", "Article", "NewsArticle"):
                    if not abstract:
                        abstract = item.get("abstract") or item.get("description")
                    if not keywords and item.get("keywords"):
                        kw = item.get("keywords")
                        keywords = [k.strip() for k in re.split(r"[,;]", kw) if k.strip()]
        except Exception:
            pass

    # Abstract quality validation
    if abstract:
        a = abstract.strip()
        if not (80 <= len(a) <= 5000):
            abstract = None
        elif re.search(r'(?i)^\s*(copyright|doi\s*[:/]|ar[xX]iv|received|accepted|published'
                       r'|conference|proceedings|volume\s+\d+|pages?\s+\d+|isbn|index terms'
                       r'|permission|\(c\)|all rights reserved)', a[:300]):
            abstract = None
        elif len(re.findall(r'\w{4,}', a)) < 15:
            abstract = None

    return abstract if abstract else None, keywords


def _tryRegexAbstractKeywords(html: str) -> Tuple[Optional[str], List[str]]:
    """Use regex to extract abstract and keywords from HTML."""
    # Preprocess: decode HTML entities
    for old, new in [
        ('&#39;', "'"), ('&#34;', '"'), ('&amp;', '&'), ('&lt;', '<'), ('&gt;', '>'),
        ('&nbsp;', ' '), ('&quot;', '"'), ('&apos;', "'"),
    ]:
        html = html.replace(old, new)

    # Abstract section names
    abs_w = [
        'abstract', 'summary', 'abstracts',
        '1.?abstract', 'abstract.?1',
        'introduction.and.abstract',
    ]
    abs_w = [w.replace('.', '\\.') for w in abs_w]

    end_w = [
        '1.introduction', 'background', 'related.work',
        'methodology', 'preliminary', 'overview',
        'figure.\\d', 'fig.\\d', 'table.\\d',
        'references', 'bibliography', 'acknowledgements',
        'appendix', 'supplementary',
    ]
    end_w = [w.replace('.', '\\.') for w in end_w]

    label_re = '|'.join(abs_w)
    end_re = '|'.join(end_w)

    abs_pat = re.compile(
        r'(?is)(?:^|\\n)\\s*(' + label_re + r')\\s*[:.\\-—–]?\\s*'
        r'(.{60,2500}?)'
        r'(?=\\n\\s*(?:' + end_re + r')\\b|\\Z)',
        re.IGNORECASE,
    )
    m = abs_pat.search(html)
    abstract = None
    if m:
        raw = m.group(2)
        raw = re.sub(r'(?m)^\\s*\\d+\\s+(?=\\S)', '', raw)
        raw = re.sub(r'(?m)^\\s*[-*#\\u2022\\u2023\\u25e6]\\s+', '', raw)
        raw = re.sub(r'\\s{2,}', ' ', raw)
        raw = re.sub(r'\\n{3,}', '\\n\\n', raw).strip()
        if len(raw) > 80 and re.search(r'[a-zA-Z]{5,}', raw):
            abstract = raw

    # Keywords
    kw_w = [
        'keywords?', 'key.words?', 'index.terms?',
        'JEL.Classification', 'ACM.CSS',
    ]
    kw_w = [w.replace('.', '\\.') for w in kw_w]
    kw_label_re = '|'.join(kw_w)
    kw_pat = re.compile(
        r'(?im)(?:^|\\n)\\s*(' + kw_label_re + r')\\s*[:.\\-—–]\\s*([^\\n]{10,600})',
    )
    km = kw_pat.search(html)
    keywords = []
    if km:
        raw_kw = km.group(2)
        for k in re.split(r'[,;、，；\\n]', raw_kw):
            k = k.strip()
            k = re.sub(r'^[\\-:_.\\s]+|[\\s]+$', '', k)
            if 2 <= len(k) <= 80 and re.search(r'[a-zA-Z一-鿿]', k):
                keywords.append(k)

    return (abstract if abstract else None), keywords


# ============ Platform-specific extractors ============

def arxiv_abstract_from_url(url: str) -> Tuple[Optional[str], List[str]]:
    """Extract abstract from arXiv page."""
    html = fetch_html(url)
    if not html:
        return None, []
    soup = BeautifulSoup(html, "html.parser")
    bq = soup.find("blockquote", class_="abstract")
    if bq:
        raw = bq.get_text(" ", strip=True)
        raw = re.sub(r"^Abstract:\s*", "", raw, flags=re.IGNORECASE).strip()
        if len(raw) > 30:
            return raw, []
    return None, []


def neurips_abstract_from_url(url: str) -> Tuple[Optional[str], List[str]]:
    """Extract abstract from NeurIPS proceedings page."""
    html = fetch_html(url)
    if not html:
        return None, []
    abstract, keywords = _tryStructuredMeta(html)
    if not abstract:
        abstract, keywords = _tryRegexAbstractKeywords(html)
    return abstract if abstract else None, keywords


def mlr_abstract_from_url(url: str) -> Tuple[Optional[str], List[str]]:
    """Extract abstract from MLR proceedings page."""
    html = fetch_html(url)
    if not html:
        return None, []
    abstract, keywords = _tryStructuredMeta(html)
    if not abstract:
        abstract, keywords = _tryRegexAbstractKeywords(html)
    return abstract if abstract else None, keywords


def doi_abstract_from_url(url: str) -> Tuple[Optional[str], List[str]]:
    """Extract abstract from DOI landing page."""
    html = fetch_html(url)
    if not html:
        return None, []
    abstract, keywords = _tryStructuredMeta(html)
    if not abstract:
        abstract, keywords = _tryRegexAbstractKeywords(html)
    return abstract if abstract else None, keywords


def openreview_abstract_from_url(url: str) -> Tuple[Optional[str], List[str]]:
    """Extract abstract from OpenReview page."""
    html = fetch_html(url)
    if not html:
        return None, []

    # 1. Structured metadata
    abstract, keywords = _tryStructuredMeta(html)
    if abstract:
        return abstract, keywords

    # 2. citation_abstract meta tag
    for pattern in [
        r'content="([^"\]{80,})"',
        r"content='([^'\]{80,})'",
        r'content="([^"\]{80,})"[^>]+name="citation_abstract"',
        r"content='([^'\]{80,})'[^>]+name='citation_abstract'",
    ]:
        m = re.search(pattern, html, re.IGNORECASE)
        if m:
            return m.group(1).strip(), []

    # 3. ICLR notes div format
    m5 = re.search(
        r'<div[^>]+class="[^"]*note[^"]*"[^>]*>.*?<h4[^>]*>[^<]*(?:abstract|summary)[^<]*</h4>.*?<p>(.{100,2000}?)</p>',
        html, re.IGNORECASE | re.DOTALL
    )
    if m5:
        txt = m5.group(1).strip()
        if len(txt) > 80:
            return txt, []

    # 4. keywords meta tag
    kws = []
    mk = re.search(r'name="keywords"[^>]+content="([^"]{10,})"', html, re.IGNORECASE)
    if not mk:
        mk = re.search(r"name='keywords'[^>]+content='([^']{10,})'", html, re.IGNORECASE)
    if mk:
        raw_kw = mk.group(1).strip()
        for k in re.split(r'[,;\n]', raw_kw):
            k = k.strip()
            if 2 <= len(k) <= 80:
                kws.append(k)

    # 5. Regex fallback
    abstract_regex, kw_regex = _tryRegexAbstractKeywords(html)
    if not kws and kw_regex:
        kws = kw_regex
    return abstract_regex, kws


# ============ Abstract extraction main function ============

def _extract_from_url(url: str) -> Tuple[Optional[str], List[str]]:
    """Dispatch to appropriate extractor based on URL type."""
    url_lower = url.lower()

    if 'proceedings.neurips.cc' in url_lower:
        return neurips_abstract_from_url(url)
    elif 'proceedings.mlr.press' in url_lower:
        return mlr_abstract_from_url(url)
    elif 'openreview.net' in url_lower:
        return openreview_abstract_from_url(url)
    elif 'doi.org' in url_lower:
        return doi_abstract_from_url(url)
    elif 'arxiv.org' in url_lower:
        return arxiv_abstract_from_url(url)

    return None, []


def _extract_keywords_from_text(text: str) -> List[str]:
    """Extract keywords from text."""
    keyword_patterns = [
        re.compile(r'(?im)^keywords?\s*[:\n]\s*(.*?)(?=\n\s*(?:\w|1\s*introduction|introduction|$))'),
        re.compile(r'(?im)^\d+\s*keywords?\s*[:\n]\s*(.*?)(?=\n\s*(?:\w|1\s*introduction|introduction|$))'),
        re.compile(r'(?im)\nkeywords?\s*[:\n]\s*(.*?)(?=\n\s*(?:\w|1\s*introduction|introduction|$))'),
        re.compile(r'(?im)(?:^|\n)\s*subject\s+keywords?\s*[:\n]\s*(.*?)(?=\n\s*(?:\w|1\s*introduction|introduction|$))'),
    ]

    for pat in keyword_patterns:
        m = pat.search(text)
        if m:
            kw_text = m.group(1).strip()
            return [k.strip().rstrip(",;.") for k in re.split(r'[,;]+', kw_text) if k.strip()]

    return []


def build_abstract_keywords(
    title: str,
    year: Optional[int],
    ee_list: List[str],
    doi_hint: Optional[str] = None,
    openalex_work: Optional[dict] = None,
) -> Tuple[Optional[str], List[str]]:
    """
    Build abstract and keywords.

    Args:
        title: Paper title
        year: Publication year
        ee_list: URL list
        doi_hint: DOI hint
        openalex_work: OpenAlex work object

    Returns:
        (abstract, keywords list)
    """
    # 1. Try DOI first
    if doi_hint:
        abstract, keywords = openalex_abstract_from_doi(doi_hint)
        if abstract:
            return abstract, keywords

    # 2. Iterate URL list
    for url in ee_list:
        if not url:
            continue

        abstract, keywords = _extract_from_url(url)
        if abstract:
            if not keywords:
                keywords = _extract_keywords_from_text(abstract)
            return abstract, keywords

    # 3. From OpenAlex work object
    if openalex_work:
        om = openalex_work.get('abstract_inverted_index')
        if om:
            abstract = inverted_index_to_text(om).strip() or None
            if abstract:
                return abstract, []

        bib = openalex_work.get('bibtex', '')
        if bib:
            m = re.search(r'(?is)abstract\s*=\s*[{]?(.*?)[}]', bib)
            if m:
                return m.group(1).strip(), []

    return None, []


# ============ PDF processing ============

def extract_text_from_pdf(pdf_bytes: bytes, max_pages: int = 30) -> str:
    """Extract plain text from PDF."""
    if fitz is None:
        return ""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception:
        return ""
    text_parts = []
    pages_to_read = min(len(doc), max_pages)
    for i in range(pages_to_read):
        try:
            text_parts.append(doc[i].get_text("text"))
        except Exception:
            continue
    full_text = "\n".join(text_parts)
    full_text = full_text.replace("\r\n", "\n")
    full_text = re.sub(r"[ \t]+", " ", full_text)
    return full_text


def extract_conclusion_from_pdf(pdf_bytes: bytes) -> str:
    """Extract conclusion from PDF."""
    if fitz is None:
        return "PyMuPDF not installed, cannot extract conclusion."

    full_text = extract_text_from_pdf(pdf_bytes)
    if not full_text.strip():
        return "PDF text is empty, could not extract conclusion."

    # Match conclusion section
    pattern = re.compile(
        r"(?is)(?:^|\n)\s*(?:\d+(?:\.\d+)*\s*)?"
        r"(conclusion(?:s)?|concluding remarks|discussion and conclusion|summary|summary and conclusions|final remarks|concluding summary|main results?)\s*[:\n]"
        r"(.*?)(?=\n\s*(?:references|acknowledg(?:e)?ments?|appendix|supplementary|bibliography)\b|\Z)",
        re.IGNORECASE,
    )
    m = pattern.search(full_text)
    if m:
        candidate = m.group(2)
        candidate = re.sub(r"\n{2,}", "\n", candidate).strip()
        candidate = re.sub(r"\s+", " ", candidate)
        if len(candidate) > 2000:
            candidate = candidate[:2000].rstrip() + "..."
        if len(candidate) >= 60:
            return candidate

    # Fallback: take last few paragraphs
    ref_pattern = re.compile(
        r"(?i)\n\s*(references|bibliography|acknowledg|appendix)\b",
    )
    ref_match = ref_pattern.search(full_text)
    if ref_match:
        last_part = full_text[ref_match.start():]
    else:
        last_part = full_text

    paragraphs = [p.strip() for p in last_part.split("\n") if p.strip() and len(p.strip()) > 40]
    if paragraphs:
        fallback_candidate = " ".join(paragraphs[-3:])
        fallback_candidate = re.sub(r"\s+", " ", fallback_candidate)
        if len(fallback_candidate) > 2000:
            fallback_candidate = fallback_candidate[:2000].rstrip() + "..."
        if len(fallback_candidate) >= 60:
            return f"[Extracted from final section] {fallback_candidate}"

    return "Could not identify a clear Conclusion section in the open PDF."


def extract_abstract_from_pdf(pdf_bytes: bytes) -> str:
    """Extract abstract from first few pages of PDF."""
    if fitz is None:
        return ""
    full_text = extract_text_from_pdf(pdf_bytes, max_pages=5)
    if not full_text.strip():
        return ""
    pattern = re.compile(
        r"(?is)(?:^|\n)\s*(?:\d+\s*)?"
        r"(abstract|summary)\s*[:\n]"
        r"(.*?)(?=\n\s*(?:1\s*introduction|introduction|keywords?|\d+\s*\w)|\Z)",
        re.IGNORECASE,
    )
    m = pattern.search(full_text)
    if m:
        candidate = m.group(2).strip()
        candidate = re.sub(r"\s+", " ", candidate)
        if len(candidate) > 1500:
            candidate = candidate[:1500].rstrip() + "..."
        if len(candidate) >= 50:
            return candidate
    return ""


# ============ PDF URL retrieval ============

def extract_pdf_links_from_html(url: str, html: str) -> List[str]:
    """Extract PDF links from HTML."""
    soup = BeautifulSoup(html, "html.parser")
    urls: List[str] = []

    for attr in ("citation_pdf_url", "dc.identifier", "dc.relation"):
        tag = soup.find("meta", attrs={"name": attr})
        if tag and tag.get("content"):
            val = tag["content"].strip()
            if ".pdf" in val.lower() or "arxiv.org" in val.lower():
                urls.append(val)

    for meta in soup.find_all("meta"):
        content = (meta.get("content") or "").strip()
        if content and ".pdf" in content.lower():
            urls.append(content)

    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        if ".pdf" in href.lower():
            urls.append(href)

    if "ieeexplore.ieee.org/document/" in url:
        m = re.search(r"/document/(\d+)", url)
        if m:
            arnum = m.group(1)
            urls.append(f"https://ieeexplore.ieee.org/stampPDF/getPDF.jsp?tp=&arnumber={arnum}")

    dedup = []
    seen = set()
    for u in urls:
        if not u:
            continue
        abs_u = requests.compat.urljoin(url, u)
        if abs_u not in seen:
            seen.add(abs_u)
            dedup.append(abs_u)
    return dedup


def search_semantic_scholar_pdf_by_doi(doi: str) -> Optional[str]:
    """Get PDF link from Semantic Scholar via DOI."""
    doi = normalize_doi(doi)
    if not doi:
        return None
    try:
        resp = request_get(
            f"{S2_BASE}/DOI:{doi}",
            params={"fields": "openAccessPdf,url,title"},
            proxies=REQUEST_PROXIES,
            headers=DEFAULT_HEADERS,
        )
        resp.raise_for_status()
        obj = resp.json()
        pdf = (obj.get("openAccessPdf") or {}).get("url")
        return pdf.strip() if pdf else None
    except Exception:
        return None


def search_arxiv_pdf_by_title(title: str) -> Optional[str]:
    """Search for PDF on arXiv by title."""
    if not title:
        return None
    try:
        resp = request_get(
            ARXIV_API,
            params={"search_query": f'ti:"{title}"', "start": 0, "max_results": 3},
            proxies=REQUEST_PROXIES,
            headers=DEFAULT_HEADERS,
        )
        resp.raise_for_status()
        root = ET.fromstring(resp.text)
    except Exception:
        return None

    ns = {"a": "http://www.w3.org/2005/Atom"}
    best_pdf = None
    best_score = 0.0
    q = normalize_text(title)

    for entry in root.findall("a:entry", ns):
        e_title = (entry.findtext("a:title", default="", namespaces=ns) or "").strip()
        if not e_title:
            continue
        score = SequenceMatcher(None, q, normalize_text(e_title)).ratio()
        pdf_url = None
        for link in entry.findall("a:link", ns):
            if (link.attrib.get("title") or "").lower() == "pdf":
                pdf_url = link.attrib.get("href")
                break
        if score > best_score and pdf_url:
            best_score = score
            best_pdf = pdf_url

    if best_score >= 0.72:
        return best_pdf
    return None


def get_open_access_pdf_url(work: dict, title: str, preferred_urls: Optional[List[str]] = None) -> Optional[str]:
    """Get open access PDF URL."""
    candidates: List[str] = []
    if preferred_urls:
        candidates.extend([u for u in preferred_urls if u])

    for key in ("best_oa_location", "primary_location"):
        loc = work.get(key) or {}
        pdf_url = loc.get("pdf_url")
        if pdf_url:
            candidates.append(pdf_url)
        landing = loc.get("landing_page_url")
        if landing:
            candidates.append(landing)
    oa = work.get("open_access") or {}
    oa_url = oa.get("oa_url")
    if oa_url:
        candidates.append(oa_url)

    ids = work.get("ids") or {}
    doi_val = ids.get("doi") or work.get("doi") or ""
    arxiv_val = ids.get("arxiv") or ""
    if doi_val:
        candidates.append(doi_val)
        doi_norm = normalize_doi(doi_val)
        if doi_norm:
            candidates.append(f"https://doi.org/{doi_norm}")
            s2_pdf = search_semantic_scholar_pdf_by_doi(doi_norm)
            if s2_pdf:
                candidates.append(s2_pdf)
    if arxiv_val:
        arx = extract_arxiv_id(arxiv_val)
        if arx:
            candidates.append(f"https://arxiv.org/pdf/{arx}.pdf")
            candidates.append(f"https://arxiv.org/abs/{arx}")

    arxiv_pdf = search_arxiv_pdf_by_title(title)
    if arxiv_pdf:
        candidates.append(arxiv_pdf)

    visited = set()
    for c in candidates:
        if not c or c in visited:
            continue
        visited.add(c)
        low = c.lower()
        if low.endswith(".pdf") or "arxiv.org/pdf/" in low:
            return c
        html = fetch_html(c)
        if not html:
            continue
        for u in extract_pdf_links_from_html(c, html):
            if u.lower().endswith(".pdf") or "stampPDF/getPDF.jsp" in u:
                return u
    return None


def fetch_conclusion_from_work(
    work: Optional[dict],
    title: str,
    preferred_urls: Optional[List[str]] = None,
    doi_hint: Optional[str] = None,
) -> str:
    """Fetch conclusion from paper work object."""
    if not work:
        work = {
            "best_oa_location": {},
            "primary_location": {},
            "open_access": {},
            "doi": doi_hint,
            "ids": {"doi": doi_hint} if doi_hint else {},
        }

    pdf_url = get_open_access_pdf_url(work, title, preferred_urls=preferred_urls)
    if not pdf_url:
        return "No open PDF link available for this paper."

    try:
        resp = request_get(
            pdf_url,
            proxies=REQUEST_PROXIES,
            headers=DEFAULT_HEADERS,
            allow_redirects=True,
        )
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "").lower()
        if "pdf" not in content_type and not pdf_url.lower().endswith(".pdf"):
            return "Open link is not a PDF, cannot extract conclusion."
        return extract_conclusion_from_pdf(resp.content)
    except Exception:
        return "Failed to download open PDF."


# ============ Paper result construction ============

def build_paper_result(author_name: str, pub: dict) -> Optional[PaperResult]:
    """Build PaperResult from Google Scholar paper."""
    bib = pub.get("bib", {})
    title = bib.get("title", "").strip()
    if not title:
        return None

    try:
        year = int(bib.get("pub_year", 0))
    except Exception:
        year = None

    authors_raw = bib.get("author", "")
    author_position = get_author_position(author_name, authors_raw)
    is_first_author = author_position == 1
    top_three = is_top_n_author(author_name, authors_raw)

    work = search_openalex_work(title, year)
    corresponding = is_corresponding_author(author_name, work)

    if not (top_three or corresponding):
        return None

    abstract = ""
    keywords: List[str] = []
    source_url = None

    if work:
        abstract = inverted_index_to_text(work.get("abstract_inverted_index", {})).strip()
        kws = work.get("keywords") or []
        keywords = [k.get("display_name", "").strip() for k in kws if k.get("display_name")]
        source_url = work.get("id") or work.get("doi")
    cited_by_count = int((work or {}).get("cited_by_count") or 0)

    if not abstract:
        abstract = bib.get("abstract", "").strip()
        if not abstract:
            pdf_url = get_open_access_pdf_url(work, title) if work else None
            if pdf_url:
                try:
                    resp = request_get(
                        pdf_url,
                        proxies=REQUEST_PROXIES,
                        headers=DEFAULT_HEADERS,
                        allow_redirects=True,
                        timeout=15,
                    )
                    resp.raise_for_status()
                    content_type = resp.headers.get("content-type", "").lower()
                    if "pdf" in content_type or pdf_url.lower().endswith(".pdf"):
                        abstract = extract_abstract_from_pdf(resp.content)
                except Exception:
                    pass
        if not abstract:
            abstract = "Abstract not available."

    conclusion = fetch_conclusion_from_work(work, title)

    return PaperResult(
        title=xml_safe_text(title),
        year=year,
        authors=xml_safe_text(authors_raw or "Unknown"),
        author_position=author_position,
        is_first_author=is_first_author,
        is_top_three_author=top_three,
        is_corresponding_author=corresponding,
        abstract=xml_safe_text(abstract),
        keywords=[xml_safe_text(k) for k in keywords],
        conclusion=xml_safe_text(conclusion),
        source_url=xml_safe_text(source_url) if source_url else None,
        cited_by_count=cited_by_count,
    )


def build_paper_result_from_dblp(author_name: str, pub: dict) -> Optional[PaperResult]:
    """Build PaperResult from DBLP paper."""
    title = (pub.get("title") or "").strip()
    if not title:
        return None

    year = pub.get("year")
    authors_raw = pub.get("authors_raw") or "Unknown"
    author_position = get_author_position(author_name, authors_raw)
    is_first_author = author_position == 1
    top_three = is_top_n_author(author_name, authors_raw)

    work = search_openalex_work(title, year)
    corresponding = is_corresponding_author(author_name, work)
    if not (top_three or corresponding):
        return None

    ee_list = pub.get("ee_list") or []
    preferred_urls = [u for u in ee_list if u]
    if pub.get("rec_url"):
        preferred_urls.append(pub.get("rec_url"))
    source_url = preferred_urls[0] if preferred_urls else pub.get("rec_url")

    abstract, keywords = build_abstract_keywords(
        title=title,
        year=year,
        ee_list=ee_list,
        doi_hint=pub.get("doi"),
        openalex_work=work,
    )
    if not abstract:
        abstract = "Abstract not available."

    cited_by_count = int((work or {}).get("cited_by_count") or 0)
    if work and not source_url:
        source_url = work.get("id") or work.get("doi") or source_url

    conclusion = fetch_conclusion_from_work(
        work,
        title,
        preferred_urls=preferred_urls,
        doi_hint=pub.get("doi"),
    )

    return PaperResult(
        title=xml_safe_text(title),
        year=year,
        authors=xml_safe_text(authors_raw),
        author_position=author_position,
        is_first_author=is_first_author,
        is_top_three_author=top_three,
        is_corresponding_author=corresponding,
        abstract=xml_safe_text(abstract),
        keywords=[xml_safe_text(k) for k in keywords],
        conclusion=xml_safe_text(conclusion),
        source_url=xml_safe_text(source_url) if source_url else None,
        cited_by_count=cited_by_count,
    )


# ============ DBLP processing ============

def fetch_dblp_author_and_publications(dblp_url: str) -> Tuple[Optional[str], List[dict], Optional[Exception]]:
    """Fetch author and publication list from DBLP."""
    pid = extract_dblp_pid(dblp_url)
    if not pid:
        return None, [], ValueError("Could not parse PID from DBLP URL.")

    xml_url = f"https://dblp.org/pid/{pid}.xml"
    try:
        resp = request_get(xml_url, proxies=REQUEST_PROXIES, headers=DEFAULT_HEADERS)
        resp.raise_for_status()
        root = ET.fromstring(resp.text)
    except Exception as exc:
        return None, [], exc

    author_name = (root.attrib.get("name") or "").strip() or None
    pubs: List[dict] = []

    for r in root.findall("./r"):
        if not list(r):
            continue
        entry = list(r)[0]
        title = " ".join("".join(entry.find("title").itertext()).split()).strip() if entry.find("title") is not None else ""
        if not title:
            continue

        authors = []
        for a in entry.findall("author"):
            author_text = " ".join("".join(a.itertext()).split()).strip()
            if author_text:
                authors.append(author_text)
        authors_raw = ", ".join(authors)

        year_val = " ".join("".join(entry.find("year").itertext()).split()).strip() if entry.find("year") is not None else ""
        try:
            year = int(year_val) if year_val else None
        except Exception:
            year = None

        doi = " ".join("".join(entry.find("doi").itertext()).split()).strip() if entry.find("doi") is not None else ""
        ee_list = []
        for ee in entry.findall("ee"):
            ee_text = " ".join("".join(ee.itertext()).split()).strip()
            if ee_text:
                ee_list.append(ee_text)
        key = (entry.attrib.get("key") or "").strip()
        rec_url = f"https://dblp.org/rec/{key}.html" if key else None

        pubs.append({
            "title": title.rstrip("."),
            "year": year,
            "authors_raw": authors_raw,
            "doi": doi,
            "ee_list": ee_list,
            "rec_url": rec_url,
        })

    return author_name, pubs, None


# ============ Mainline graph ============

def tokenize_topic_text(text: str) -> List[str]:
    """Tokenize topic text."""
    text = normalize_text(text)
    if not text:
        return []
    english = re.findall(r"[a-z][a-z0-9+\-]{1,}", text)
    chinese = re.findall(r"[一-鿿]{2,8}", text)
    return english + chinese


def load_mainline_profile(path: str) -> Dict[str, object]:
    """Load mainline graph."""
    if not path:
        return {"topics": [], "external_evidence_count": 0}
    try:
        with open(path, "r", encoding="utf-8") as f:
            obj = json.load(f)
    except Exception:
        return {"topics": [], "external_evidence_count": 0}

    topics: List[dict] = []
    nodes = obj.get("nodes") or []
    for node in nodes:
        name = (node.get("topic") or node.get("name") or "").strip()
        if not name:
            continue
        aliases = node.get("aliases") or []
        all_aliases = [a for a in aliases if isinstance(a, str)] + [name]
        topic_tokens = set()
        for alias in all_aliases:
            topic_tokens.update(tokenize_topic_text(alias))
        topics.append({
            "name": name,
            "tokens": topic_tokens,
            "raw_aliases": all_aliases,
            "first_year": node.get("first_year"),
            "last_year": node.get("last_year"),
            "evidence_count": int(node.get("evidence_count") or 0),
            "weight": float(node.get("weight") or 0.0),
        })
    external_evidence_count = len([s for s in (obj.get("sources") or []) if isinstance(s, dict)])
    return {"topics": topics, "external_evidence_count": external_evidence_count}


def build_mainline_profile_from_papers(papers: List[PaperResult], top_k: int = 8) -> Dict[str, object]:
    """Build mainline profile from paper list."""
    counter = Counter()
    for p in papers:
        for kw in p.keywords or []:
            for token in tokenize_topic_text(kw):
                if len(token) >= 3:
                    counter[token] += 1

    topics = []
    for token, count in counter.most_common(top_k):
        topics.append({
            "name": token,
            "tokens": {token},
            "raw_aliases": [token],
            "first_year": None,
            "last_year": None,
            "evidence_count": int(count),
            "weight": min(1.0, count / 5.0),
        })
    return {"topics": topics, "external_evidence_count": 0}


def score_topic_match(text: str, topic: dict) -> float:
    """Calculate text-topic match score."""
    raw_aliases = topic.get("raw_aliases") or []
    n_text = normalize_text(text)
    if not n_text:
        return 0.0

    substring_bonus = 0.0
    for alias in raw_aliases:
        a = normalize_text(alias)
        if a and a in n_text:
            substring_bonus = 1.0
            break

    paper_tokens = set(tokenize_topic_text(n_text))
    topic_tokens = topic.get("tokens") or set()
    if not topic_tokens:
        return substring_bonus * 0.6

    overlap = len(paper_tokens & topic_tokens)
    ratio = overlap / max(len(topic_tokens), 1)
    return min(1.0, max(ratio, substring_bonus * 0.75))


# ============ Scoring ============

def compute_authorship_weight(paper: PaperResult) -> float:
    """Compute authorship weight."""
    if paper.is_corresponding_author and paper.is_first_author:
        return 0.60
    if paper.is_corresponding_author or paper.is_first_author:
        return 0.45
    if 1 <= paper.author_position <= 3:
        return 0.25
    if paper.author_position > 0:
        return 0.10
    return 0.0


def compute_citation_impact(paper: PaperResult) -> float:
    """Compute citation impact."""
    if paper.cited_by_count <= 0:
        return 0.0
    current_year = time.localtime().tm_year
    pub_year = paper.year or current_year
    age = max(1, current_year - pub_year + 1)
    cited_per_year = max(0.0, float(paper.cited_by_count) / float(age))
    return min(1.0, math.log1p(cited_per_year) / math.log1p(30.0))


def compute_representativeness(paper: PaperResult, mainline_profile: Dict[str, object]) -> None:
    """Compute paper representativeness score."""
    authorship_weight = compute_authorship_weight(paper)

    evidence_text = "\n".join([paper.title, paper.abstract, " ".join(paper.keywords), paper.conclusion])
    topics = mainline_profile.get("topics") or []
    best_topic_name = ""
    best_topic_score = 0.0
    best_topic_obj = None
    for topic in topics:
        s = score_topic_match(evidence_text, topic)
        weighted = min(1.0, 0.8 * s + 0.2 * float(topic.get("weight") or 0.0))
        if weighted > best_topic_score:
            best_topic_score = weighted
            best_topic_name = topic.get("name") or ""
            best_topic_obj = topic

    topic_persistence = 0.25
    if best_topic_obj:
        first_year = best_topic_obj.get("first_year")
        last_year = best_topic_obj.get("last_year")
        if isinstance(first_year, int) and isinstance(last_year, int) and last_year >= first_year:
            span = last_year - first_year + 1
            topic_persistence = min(1.0, span / 5.0)
        else:
            topic_persistence = 0.55 if best_topic_score > 0.55 else 0.35
    citation_impact = compute_citation_impact(paper)
    external_evidence_bonus = min(1.0, float(mainline_profile.get("external_evidence_count") or 0) / 6.0)

    score = (
        0.30 * authorship_weight
        + 0.30 * best_topic_score
        + 0.15 * topic_persistence
        + 0.15 * citation_impact
        + 0.10 * external_evidence_bonus
    )

    if score >= 0.60:
        bucket = "core"
    elif score >= 0.40:
        bucket = "peripheral"
    else:
        bucket = "student-led"

    paper.authorship_weight = round(authorship_weight, 4)
    paper.mainline_relevance = round(best_topic_score, 4)
    paper.topic_persistence = round(topic_persistence, 4)
    paper.citation_impact = round(citation_impact, 4)
    paper.external_evidence_bonus = round(external_evidence_bonus, 4)
    paper.self_mention_bonus = paper.external_evidence_bonus
    paper.representativeness_score = round(min(1.0, max(0.0, score)), 4)
    paper.bucket = bucket
    paper.matched_topic = best_topic_name


# ============ Export functions ============

def export_scored_json(author_name: str, papers: List[PaperResult], output_path: str) -> None:
    """Export scored papers as JSON."""
    obj = {
        "author_name": author_name,
        "paper_count": len(papers),
        "papers": [asdict(p) for p in papers],
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)


def apply_document_fonts(doc: Document) -> None:
    """Apply document fonts."""
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def export_to_word(author_name: str, papers: List[PaperResult], output_path: str, max_papers: int) -> None:
    """Export to Word document."""
    doc = Document()
    apply_document_fonts(doc)
    doc.add_heading(xml_safe_text(f"DBLP Literature Review - {author_name}"), level=1)
    doc.add_paragraph(xml_safe_text(f"Total {len(papers)} papers reviewed (max {max_papers})."))
    doc.add_paragraph(
        xml_safe_text("Note: Top-3 authorship determined by author order; corresponding author based on OpenAlex is_corresponding field (if available).")
    )

    for i, p in enumerate(papers, 1):
        doc.add_heading(xml_safe_text(f"{i}. {p.title}"), level=2)
        role = []
        if p.is_top_three_author:
            role.append("Top-3 Author")
        if p.is_corresponding_author:
            role.append("Corresponding Author")
        role_text = ", ".join(role) if role else "Not determined"

        doc.add_paragraph(xml_safe_text(f"Year: {p.year or 'Unknown'}"))
        doc.add_paragraph(xml_safe_text(f"Authors: {p.authors}"))
        doc.add_paragraph(xml_safe_text(f"Author Role: {role_text}"))
        doc.add_paragraph(
            xml_safe_text(
                f"Representativeness Score: {p.representativeness_score:.3f} | Bucket: {p.bucket} | Matched Topic: {p.matched_topic or 'None'}"
            )
        )
        doc.add_paragraph(
            xml_safe_text(
                "Score Breakdown: "
                f"Authorship={p.authorship_weight:.3f}, "
                f"Mainline Relevance={p.mainline_relevance:.3f}, "
                f"Topic Persistence={p.topic_persistence:.3f}, "
                f"Citation Impact={p.citation_impact:.3f}, "
                f"External Evidence={p.external_evidence_bonus:.3f}"
            )
        )

        if p.source_url:
            doc.add_paragraph(xml_safe_text(f"Metadata Source: {p.source_url}"))

        doc.add_paragraph(xml_safe_text("Abstract:"))
        doc.add_paragraph(xml_safe_text(p.abstract))

        doc.add_paragraph(xml_safe_text("Keywords:"))
        doc.add_paragraph(xml_safe_text("; ".join(p.keywords) if p.keywords else "No keywords available."))

        doc.add_paragraph(xml_safe_text("Conclusion:"))
        doc.add_paragraph(xml_safe_text(p.conclusion))

    doc.save(output_path)


# ============ Main function ============

def main():
    """Main function."""
    parser = argparse.ArgumentParser(
        description="Fetch top-3/corresponding author papers from DBLP and export as Word."
    )
    parser.add_argument(
        "--dblp-url",
        type=str,
        help="Researcher DBLP homepage URL, e.g. https://dblp.org/pid/xx/xxxx.html",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="dblp_author_report.docx",
        help="Output Word filename (default: dblp_author_report.docx)",
    )
    parser.add_argument(
        "--max-papers",
        type=int,
        default=MAX_PAPERS,
        help=f"Max papers to process (default: {MAX_PAPERS})",
    )
    parser.add_argument("--http-proxy", type=str, default="", help="HTTP proxy, e.g. http://127.0.0.1:7890")
    parser.add_argument("--https-proxy", type=str, default="", help="HTTPS proxy, e.g. http://127.0.0.1:7890")
    parser.add_argument(
        "--mainline-graph",
        type=str,
        default="",
        help="Optional: mainline graph JSON (mainline_graph.json) for paper attribution weighting",
    )
    parser.add_argument(
        "--json-output",
        type=str,
        default="",
        help="Optional: output scored JSON file; defaults to same name as Word with .json extension",
    )
    args = parser.parse_args()

    output_path = (args.output or "").strip()
    if not output_path:
        output_path = "scholar_author_report.docx"
    max_papers = max(1, int(args.max_papers or MAX_PAPERS))

    global REQUEST_PROXIES
    REQUEST_PROXIES = configure_proxies(args.http_proxy, args.https_proxy)
    mainline_profile = load_mainline_profile((args.mainline_graph or "").strip())

    dblp_url = (args.dblp_url or "").strip()
    if not dblp_url:
        dblp_url = input("Enter researcher DBLP homepage URL: ").strip()
    if not dblp_url:
        print("No DBLP URL provided, exiting.")
        return

    print("[1/4] Fetching author and publications from DBLP")
    author_name, pubs, dblp_error = fetch_dblp_author_and_publications(dblp_url)
    if dblp_error:
        print(f"DBLP fetch failed: {type(dblp_error).__name__}: {dblp_error}")
        print("Try adding: --https-proxy http://127.0.0.1:7890")
        return

    author_name = author_name or "Unknown Author"
    if not pubs:
        print("No publications found on this DBLP author page.")
        return

    print(f"[2/4] Found {len(pubs)} papers, filtering top-3/corresponding author papers (max {max_papers})")

    results: List[PaperResult] = []
    seen_title_keys = set()
    for idx, p in enumerate(pubs, 1):
        if len(results) >= max_papers:
            break
        try:
            p_title = (p.get("title") or "").strip()
            title_key = canonical_title_key(p_title)
            if title_key and title_key in seen_title_keys:
                print(f"  - Skipping duplicate [{idx}/{len(pubs)}]: {p_title[:80]}")
                continue
            if p_title:
                print(f"  - Processing [{idx}/{len(pubs)}]: {p_title[:80]}")
            else:
                print(f"  - Processing [{idx}/{len(pubs)}]")
            item = build_paper_result_from_dblp(author_name, p)
            if item:
                item_key = canonical_title_key(item.title)
                if item_key and item_key in seen_title_keys:
                    print(f"  - Skipping duplicate result: {item.title}")
                    continue
                results.append(item)
                if item_key:
                    seen_title_keys.add(item_key)
                print(f"  - [{len(results)}/{max_papers}] Added: {item.title}")
        except Exception:
            continue
        time.sleep(2.0)  # Rate limiting for DBLP

    if not results:
        print("No papers matching criteria found (may lack top-3 author order/corresponding author metadata).")
        return

    if not (mainline_profile.get("topics") or []):
        mainline_profile = build_mainline_profile_from_papers(results)

    for p in results:
        compute_representativeness(p, mainline_profile)
    results.sort(key=lambda x: x.representativeness_score, reverse=True)

    print("[3/4] Generating Word document")
    export_to_word(author_name, results, output_path, max_papers)

    if args.json_output.strip():
        json_output_path = args.json_output.strip()
    else:
        base, _ = os.path.splitext(output_path)
        json_output_path = f"{base}.json"
    export_scored_json(author_name, results, json_output_path)

    print(f"[4/4] Done, output: {output_path}")
    print(f"      Scored JSON: {json_output_path}")


# ============ Callable interface for pipeline ============

def run_single(scholar_name: str, output_dir: str, mainline_file: str = None,
               dblp_url: str = "", max_papers: int = MAX_PAPERS) -> str:
    """Score and rank papers for a single scholar.

    Args:
        scholar_name: Full name of the scholar.
        output_dir: Output directory path.
        mainline_file: Path to mainline graph JSON.
        dblp_url: DBLP URL (searched from CSV if empty).
        max_papers: Maximum papers to process.

    Returns:
        Path to the scored JSON file.
    """
    import csv as csv_mod

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    json_path = out / "paper_digest_scored.json"
    docx_path = out / "scholar_report.docx"

    # Get dblp_url from CSV if not provided
    if not dblp_url:
        csv_path = out / "scholars.csv"
        if csv_path.exists():
            with open(csv_path, "r", encoding="utf-8-sig", newline="") as f:
                reader = csv_mod.DictReader(f)
                for row in reader:
                    if (row.get("author") or "").strip().lower() == scholar_name.lower():
                        dblp_url = (row.get("dblp_url") or "").strip()
                        break

    if not dblp_url:
        raise RuntimeError(f"No DBLP URL found for '{scholar_name}'.")

    mainline_profile = load_mainline_profile(mainline_file or "")

    author_name, pubs, dblp_error = fetch_dblp_author_and_publications(dblp_url)
    if dblp_error:
        raise RuntimeError(f"DBLP fetch failed: {dblp_error}")

    author_name = author_name or scholar_name
    if not pubs:
        raise RuntimeError(f"No publications found for '{scholar_name}'.")

    results: List[PaperResult] = []
    seen_title_keys = set()
    for p in pubs:
        if len(results) >= max_papers:
            break
        try:
            p_title = (p.get("title") or "").strip()
            title_key = canonical_title_key(p_title)
            if title_key and title_key in seen_title_keys:
                continue
            item = build_paper_result_from_dblp(author_name, p)
            if item:
                item_key = canonical_title_key(item.title)
                if item_key and item_key in seen_title_keys:
                    continue
                results.append(item)
                if item_key:
                    seen_title_keys.add(item_key)
        except Exception:
            continue
        time.sleep(2.0)

    if not results:
        raise RuntimeError(f"No qualifying papers found for '{scholar_name}'.")

    if not (mainline_profile.get("topics") or []):
        mainline_profile = build_mainline_profile_from_papers(results)

    for p in results:
        compute_representativeness(p, mainline_profile)
    results.sort(key=lambda x: x.representativeness_score, reverse=True)

    export_to_word(author_name, results, str(docx_path), max_papers)
    export_scored_json(author_name, results, str(json_path))

    print(f"[scholar_author_digest] {len(results)} papers scored -> {json_path}")
    return str(json_path)


if __name__ == "__main__":
    main()
